import uuid
from pathlib import Path

from src.rag.models import Document
from src.utils.logger import get_logger

logger = get_logger(__name__)

MIN_IMAGE_BYTES = 10 * 1024  # Skip images smaller than 10KB (icons, deco)


class MarkdownLoader:
    def load(self, file_path: str) -> list[Document]:
        path = Path(file_path)
        content = path.read_text(encoding="utf-8")
        doc_id = str(uuid.uuid4())
        logger.info("Loaded markdown: %s (%d chars)", path.name, len(content))
        return [
            Document(
                content=content,
                metadata={
                    "source": str(path.resolve()),
                    "file_name": path.name,
                    "file_type": "markdown",
                    "title": path.stem,
                },
                doc_id=doc_id,
            )
        ]


class PDFLoader:
    def load(self, file_path: str) -> list[Document]:
        import fitz  # PyMuPDF

        path = Path(file_path)
        docs: list[Document] = []
        doc = fitz.open(str(path))
        for page_num, page in enumerate(doc):
            text = page.get_text()
            if not text.strip():
                continue
            doc_id = str(uuid.uuid4())
            docs.append(
                Document(
                    content=text,
                    metadata={
                        "source": str(path.resolve()),
                        "file_name": path.name,
                        "file_type": "pdf",
                        "page_num": page_num + 1,
                        "title": path.stem,
                    },
                    doc_id=doc_id,
                )
            )
        doc.close()
        logger.info("Loaded PDF: %s (%d pages)", path.name, len(docs))
        return docs

    @staticmethod
    def extract_images(file_path: str, output_dir: str) -> list[str]:
        """Extract embedded images from a PDF file.

        Saves images as PNG to output_dir. Skips images smaller than MIN_IMAGE_BYTES.
        Returns a list of saved image paths.
        """
        import fitz

        path = Path(file_path)
        out = Path(output_dir)
        out.mkdir(parents=True, exist_ok=True)

        saved: list[str] = []
        doc = fitz.open(str(path))
        stem = path.stem.replace(" ", "_")

        for page_num, page in enumerate(doc):
            images = page.get_images(full=True)
            for idx, img in enumerate(images):
                xref = img[0]
                base_image = doc.extract_image(xref)
                image_bytes = base_image.get("image")
                if not image_bytes or len(image_bytes) < MIN_IMAGE_BYTES:
                    continue
                ext = base_image.get("ext", "png")
                img_name = f"{stem}_p{page_num + 1}_{idx}.{ext}"
                img_path = out / img_name
                img_path.write_bytes(image_bytes)
                saved.append(str(img_path.resolve()))
                logger.debug("Extracted image: %s (%d bytes)", img_name, len(image_bytes))

        doc.close()
        logger.info("Extracted %d images from PDF: %s", len(saved), path.name)
        return saved


class DOCXLoader:
    def load(self, file_path: str) -> list[Document]:
        from docx import Document as DocxDocument

        path = Path(file_path)
        docx = DocxDocument(str(path))
        paragraphs: list[str] = []
        for para in docx.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        content = "\n\n".join(paragraphs)
        doc_id = str(uuid.uuid4())
        logger.info("Loaded DOCX: %s (%d paragraphs)", path.name, len(paragraphs))
        return [
            Document(
                content=content,
                metadata={
                    "source": str(path.resolve()),
                    "file_name": path.name,
                    "file_type": "docx",
                    "title": path.stem,
                },
                doc_id=doc_id,
            )
        ]


LOADER_MAP = {
    ".md": MarkdownLoader,
    ".markdown": MarkdownLoader,
    ".pdf": PDFLoader,
    ".docx": DOCXLoader,
}


def load_document(file_path: str) -> list[Document]:
    path = Path(file_path)
    ext = path.suffix.lower()
    loader_cls = LOADER_MAP.get(ext)
    if loader_cls is None:
        raise ValueError(f"Unsupported file type: {ext}")
    return loader_cls().load(str(path))


def extract_document_images(file_path: str, output_dir: str) -> list[str]:
    """Extract embedded images from a document file.

    Currently supports PDF. Returns list of saved image paths.
    """
    path = Path(file_path)
    ext = path.suffix.lower()
    if ext == ".pdf":
        return PDFLoader.extract_images(file_path, output_dir)
    # DOCX image extraction not yet supported
    logger.debug("Image extraction not supported for: %s", ext)
    return []


def load_documents_from_directory(directory: str) -> list[Document]:
    dir_path = Path(directory)
    all_docs: list[Document] = []
    for ext, loader_cls in LOADER_MAP.items():
        for file_path in dir_path.glob(f"*{ext}"):
            try:
                docs = loader_cls().load(str(file_path))
                all_docs.extend(docs)
            except Exception as e:
                logger.warning("Failed to load %s: %s", file_path.name, e)
    logger.info("Loaded %d documents from %s", len(all_docs), directory)
    return all_docs
